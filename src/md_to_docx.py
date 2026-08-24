"""Reverse conversion: Markdown (.md) -> Word (.docx).

The forward pipeline renders mammoth's HTML as Markdown; this module walks the
other way, parsing the same Markdown subset - headings, emphasis, links,
images, lists, tables, block quotes, fenced code and rules - and writing it
with python-docx.

Round-tripping is deliberate. Paragraph styles are named after the ones
`DOCX_STYLE_MAP` already understands (`Code` for code blocks, `Quote` for
block quotes), so converting the result back to Markdown reproduces the
document it came from.
"""

from __future__ import annotations

import re
from dataclasses import dataclass, field, replace as _replace
from itertools import groupby
from pathlib import Path
from urllib.parse import unquote, urlparse

from docx import Document
from docx.enum.style import WD_STYLE_TYPE
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK
from docx.opc.constants import RELATIONSHIP_TYPE as RT
from docx.oxml import OxmlElement, parse_xml
from docx.oxml.ns import nsdecls, qn
from docx.shared import Emu, Pt, RGBColor

MARKDOWN_EXTENSIONS = {".md", ".markdown", ".mdown", ".mkd"}

# Word refuses control characters; they only ever arrive by accident.
_CONTROL_CHARS = re.compile(r"[\x00-\x08\x0b\x0c\x0e-\x1f]")

# One twip is 1/1440 inch. Word's own list templates step by half an inch.
_LEVEL_INDENT = 720
_HANGING_INDENT = 360
_TWIP = 635  # EMU per twip

DEFAULT_FONT = "Times New Roman"
DEFAULT_FONT_SIZE = 13.0
DEFAULT_LINE_SPACING = 1.15

# Page sizes as (width, height) in EMU, on the twip grid Word stores them on.
PAGE_SIZES = {
    "A4": (7560310, 10692130),  # 11906 x 16838 twips - 21.0 x 29.7 cm
    "Letter": (7772400, 10058400),  # 12240 x 15840 twips - 8.5 x 11 in
    "Legal": (7772400, 12801600),  # 12240 x 20160 twips - 8.5 x 14 in
}

_CODE_FONT = "Consolas"
# Code sits a little below the body text, the way Word's own Quote style does.
_CODE_SIZE_RATIO = 0.85

# Built-in styles that take their font from the theme rather than from Normal,
# so they need naming explicitly as well.
_THEMED_STYLES = (
    "Title",
    "Subtitle",
    "Quote",
    "Intense Quote",
    "List Paragraph",
    *(f"Heading {level}" for level in range(1, 10)),
)

# Word's own bullet glyphs: they live in the Symbol / Wingdings private area.
_BULLET_CHARS = [chr(0xF0B7), "o", chr(0xF0A7)]
_BULLET_FONTS = ["Symbol", "Courier New", "Wingdings"]
_ORDERED_FORMATS = ["decimal", "lowerLetter", "lowerRoman"]


# ------------------------------------------------------------------ blocks


@dataclass
class Heading:
    level: int
    text: str


@dataclass
class Paragraph:
    text: str


@dataclass
class CodeBlock:
    text: str
    language: str = ""


@dataclass
class Quote:
    blocks: list = field(default_factory=list)


@dataclass
class ListItem:
    blocks: list = field(default_factory=list)
    checked: bool | None = None


@dataclass
class ListBlock:
    ordered: bool
    start: int = 1
    items: list[ListItem] = field(default_factory=list)


@dataclass
class Table:
    header: list[str]
    rows: list[list[str]]
    alignments: list[str]


@dataclass
class Rule:
    pass


@dataclass(frozen=True)
class DocxSettings:
    """How the generated Word document should look.

    An empty `font` means "whatever the user did not choose": Times New Roman.
    """

    font: str = DEFAULT_FONT
    font_size: float = DEFAULT_FONT_SIZE
    page_size: str = "A4"
    line_spacing: float = DEFAULT_LINE_SPACING
    page_break_before_h1: bool = False
    table_of_contents: bool = False

    @property
    def body_font(self) -> str:
        return (self.font or "").strip() or DEFAULT_FONT

    @property
    def body_size(self) -> Pt:
        return Pt(self.font_size or DEFAULT_FONT_SIZE)

    @property
    def page(self) -> tuple[int, int]:
        return PAGE_SIZES.get(self.page_size, PAGE_SIZES["A4"])


_FENCE_RE = re.compile(r"^\s{0,3}(?P<fence>`{3,}|~{3,})\s*(?P<info>[^`]*)$")
_ATX_RE = re.compile(r"^\s{0,3}(?P<hashes>#{1,6})(?:\s+(?P<text>.*?))?\s*#*\s*$")
_HR_RE = re.compile(r"^\s{0,3}(?:(?:\*[ \t]*){3,}|(?:-[ \t]*){3,}|(?:_[ \t]*){3,})$")
_QUOTE_RE = re.compile(r"^\s{0,3}>[ \t]?(.*)$")
_ITEM_RE = re.compile(
    r"^(?P<indent>[ \t]*)(?P<marker>[-*+]|\d{1,9}[.)])(?P<space>[ \t]+)(?P<text>.*)$"
)
_TASK_RE = re.compile(r"^\[(?P<mark>[ xX])\][ \t]+(?P<rest>.*)$")
_TABLE_DELIM_RE = re.compile(
    r"^[ \t]*\|?[ \t]*:?-+:?[ \t]*(?:\|[ \t]*:?-+:?[ \t]*)*\|?[ \t]*$"
)
_SETEXT_H1_RE = re.compile(r"^\s{0,3}=+\s*$")
_SETEXT_H2_RE = re.compile(r"^\s{0,3}-+\s*$")
_HARD_BREAK_RE = re.compile(r"(?: {2,}|\\)$")
# The Latin face of the theme's major (heading) and minor (body) font.
_THEME_FACE_RE = re.compile(r'(<a:(?:major|minor)Font>\s*<a:latin typeface=")[^"]*')
_FRONT_MATTER_RE = re.compile(r"^\s*(?P<key>[A-Za-z_][\w -]*)\s*:\s*(?P<value>.*?)\s*$")


def parse_markdown(text: str) -> tuple[list, dict[str, str]]:
    """Split Markdown into blocks, peeling off any YAML front matter first."""
    lines = text.replace("\r\n", "\n").replace("\r", "\n").expandtabs(4).split("\n")
    if lines and lines[0].startswith("﻿"):
        lines[0] = lines[0][1:]
    metadata, lines = _split_front_matter(lines)
    return _parse_blocks(lines), metadata


def _split_front_matter(lines: list[str]) -> tuple[dict[str, str], list[str]]:
    if not lines or lines[0].strip() != "---":
        return {}, lines
    for index in range(1, len(lines)):
        if lines[index].strip() in ("---", "..."):
            metadata: dict[str, str] = {}
            for line in lines[1:index]:
                match = _FRONT_MATTER_RE.match(line)
                if match:
                    value = match.group("value").strip().strip("\"'")
                    metadata[match.group("key").strip().lower()] = value
            return metadata, lines[index + 1 :]
    return {}, lines


def _parse_blocks(lines: list[str]) -> list:
    blocks: list = []
    index = 0
    total = len(lines)

    while index < total:
        line = lines[index]

        if not line.strip():
            index += 1
            continue

        fence = _FENCE_RE.match(line)
        if fence:
            block, index = _read_fenced_code(lines, index, fence)
            blocks.append(block)
            continue

        heading = _ATX_RE.match(line)
        if heading:
            text = (heading.group("text") or "").strip()
            blocks.append(Heading(len(heading.group("hashes")), text))
            index += 1
            continue

        if _HR_RE.match(line):
            blocks.append(Rule())
            index += 1
            continue

        if _QUOTE_RE.match(line):
            block, index = _read_quote(lines, index)
            blocks.append(block)
            continue

        if _is_table_start(lines, index):
            block, index = _read_table(lines, index)
            blocks.append(block)
            continue

        if _ITEM_RE.match(line):
            block, index = _read_list(lines, index)
            blocks.append(block)
            continue

        if line.startswith("    "):
            block, index = _read_indented_code(lines, index)
            blocks.append(block)
            continue

        block, index = _read_paragraph(lines, index)
        if block is not None:
            blocks.append(block)

    return blocks


def _read_fenced_code(lines: list[str], index: int, fence) -> tuple[CodeBlock, int]:
    marker = fence.group("fence")[0] * 3
    body: list[str] = []
    index += 1
    while index < len(lines):
        stripped = lines[index].strip()
        if stripped.startswith(marker) and set(stripped) == {marker[0]}:
            index += 1
            break
        body.append(lines[index])
        index += 1
    return CodeBlock("\n".join(body), fence.group("info").strip()), index


def _read_indented_code(lines: list[str], index: int) -> tuple[CodeBlock, int]:
    body: list[str] = []
    while index < len(lines):
        line = lines[index]
        if line.startswith("    "):
            body.append(line[4:])
        elif not line.strip():
            body.append("")
        else:
            break
        index += 1
    while body and not body[-1].strip():
        body.pop()
    return CodeBlock("\n".join(body)), index


def _read_quote(lines: list[str], index: int) -> tuple[Quote, int]:
    inner: list[str] = []
    while index < len(lines):
        match = _QUOTE_RE.match(lines[index])
        if match:
            inner.append(match.group(1))
        elif lines[index].strip() and inner:
            inner.append(lines[index])  # lazy continuation
        else:
            break
        index += 1
    return Quote(_parse_blocks(inner)), index


def _is_table_start(lines: list[str], index: int) -> bool:
    if "|" not in lines[index] or index + 1 >= len(lines):
        return False
    following = lines[index + 1]
    return "|" in following and bool(_TABLE_DELIM_RE.match(following))


def _split_row(line: str) -> list[str]:
    """Split a table row on unescaped pipes, dropping the outer borders."""
    cells: list[str] = []
    current: list[str] = []
    escaped = False
    for char in line.strip():
        if escaped:
            current.append(char if char == "|" else f"\\{char}")
            escaped = False
        elif char == "\\":
            escaped = True
        elif char == "|":
            cells.append("".join(current))
            current = []
        else:
            current.append(char)
    if escaped:
        current.append("\\")
    cells.append("".join(current))

    if cells and not cells[0].strip():
        cells.pop(0)
    if cells and not cells[-1].strip():
        cells.pop()
    return [cell.strip() for cell in cells]


def _read_table(lines: list[str], index: int) -> tuple[Table, int]:
    header = _split_row(lines[index])
    alignments = []
    for cell in _split_row(lines[index + 1]):
        left, right = cell.startswith(":"), cell.endswith(":")
        if left and right:
            alignments.append("center")
        elif right:
            alignments.append("right")
        else:
            alignments.append("")
    index += 2

    rows: list[list[str]] = []
    while index < len(lines) and lines[index].strip() and "|" in lines[index]:
        rows.append(_split_row(lines[index]))
        index += 1

    width = max([len(header), len(alignments), *(len(row) for row in rows)])

    def pad(cells: list[str]) -> list[str]:
        return cells + [""] * (width - len(cells))

    return Table(pad(header), [pad(row) for row in rows], pad(alignments)), index


def _read_list(lines: list[str], index: int) -> tuple[ListBlock, int]:
    first = _ITEM_RE.match(lines[index])
    base = len(first.group("indent"))
    ordered = first.group("marker")[0] not in "-*+"
    start = int(first.group("marker")[:-1]) if ordered else 1
    block = ListBlock(ordered=ordered, start=start)

    pending: list[str] | None = None
    content_indent = 0

    def flush() -> None:
        if pending is not None:
            block.items.append(_make_item(pending))

    while index < len(lines):
        line = lines[index]
        indent = len(line) - len(line.lstrip(" "))
        item = _ITEM_RE.match(line)

        # Anything indented past the current item's text belongs to that item -
        # a nested list included, which the recursive parse picks up once the
        # lines are dedented.
        nested = pending is not None and line.strip() and indent >= content_indent

        if item and not nested and base <= indent <= base + 3:
            marker = item.group("marker")
            if (marker[0] in "-*+") == ordered:
                break  # a different marker family starts a new list
            flush()
            pending = [item.group("text")]
            content_indent = indent + len(marker) + len(item.group("space"))
            index += 1
            continue

        if pending is None:
            break

        if not line.strip():
            # A blank line keeps the list open only if indented content or
            # another item of the same family follows.
            if not _list_continues(lines, index + 1, base, ordered):
                break
            pending.append("")
            index += 1
            continue

        if indent >= content_indent:
            pending.append(line[content_indent:])
        elif not _is_block_start(line):
            pending.append(line.strip())  # lazy continuation
        else:
            break
        index += 1

    flush()
    return block, index


def _next_content(lines: list[str], index: int) -> str | None:
    for line in lines[index:]:
        if line.strip():
            return line
    return None


def _list_continues(lines: list[str], index: int, base: int, ordered: bool) -> bool:
    following = _next_content(lines, index)
    if following is None:
        return False
    indent = len(following) - len(following.lstrip(" "))
    if indent > base:
        return True
    item = _ITEM_RE.match(following)
    return bool(
        item
        and base <= indent <= base + 3
        and (item.group("marker")[0] in "-*+") != ordered
    )


def _is_block_start(line: str) -> bool:
    return bool(
        _ATX_RE.match(line)
        or _HR_RE.match(line)
        or _FENCE_RE.match(line)
        or _QUOTE_RE.match(line)
    )


def _make_item(raw: list[str]) -> ListItem:
    checked = None
    task = _TASK_RE.match(raw[0]) if raw else None
    if task:
        checked = task.group("mark").lower() == "x"
        raw = [task.group("rest"), *raw[1:]]
    return ListItem(_parse_blocks(raw), checked)


def _read_paragraph(lines: list[str], index: int) -> tuple[object | None, int]:
    collected: list[str] = []
    while index < len(lines):
        line = lines[index]
        if not line.strip():
            break
        if collected and _SETEXT_H1_RE.match(line):
            return Heading(1, _join_paragraph(collected).strip()), index + 1
        if collected and _SETEXT_H2_RE.match(line):
            return Heading(2, _join_paragraph(collected).strip()), index + 1
        if collected and (
            _ATX_RE.match(line)
            or _HR_RE.match(line)
            or _FENCE_RE.match(line)
            or _QUOTE_RE.match(line)
            or _ITEM_RE.match(line)
            or _is_table_start(lines, index)
        ):
            break
        collected.append(line)
        index += 1

    text = _join_paragraph(collected)
    return (Paragraph(text) if text.strip() else None), index


def _join_paragraph(lines: list[str]) -> str:
    """Join wrapped lines, keeping two-space and backslash hard breaks."""
    parts: list[str] = []
    for position, line in enumerate(lines):
        if position:
            parts.append("\n" if _HARD_BREAK_RE.search(lines[position - 1]) else " ")
        parts.append(_HARD_BREAK_RE.sub("", line.strip()))
    return "".join(parts)


# ------------------------------------------------------------------ inline


@dataclass(frozen=True)
class Style:
    bold: bool = False
    italic: bool = False
    code: bool = False
    strike: bool = False
    vertical: str = ""
    href: str = ""


@dataclass
class Span:
    kind: str = "text"  # text | image | break
    text: str = ""
    style: Style = Style()
    src: str = ""


_INLINE_RE = re.compile(
    r"(?P<escape>\\[\\`*_{}\[\]()#+\-.!|~<>])"
    r"|(?P<fence>`+)(?P<code>.+?)(?P=fence)"
    r"|!\[(?P<alt>(?:\\.|[^\]\\])*)\]\(\s*(?P<src>[^)]*?)\s*\)"
    r"|\[(?P<label>(?:\\.|!\[[^\]]*\]\([^)]*\)|[^\]\\])*)\]\(\s*(?P<href>[^)]*?)\s*\)"
    r"|<(?P<auto>(?:https?://|ftp://|mailto:)[^<>\s]+)>"
    r"|(?P<hardbreak><br\s*/?>)"
    r"|<(?P<vtag>sup|sub)>(?P<vtext>.*?)</(?P=vtag)>"
    r"|(?P<html></?[A-Za-z][A-Za-z0-9]*(?:\s[^<>]*)?/?>)"
    r"|\*\*\*(?P<bi>\S(?:.*?\S)?)\*\*\*"
    r"|\*\*(?P<bold>\S(?:.*?\S)?)\*\*"
    r"|(?<![A-Za-z0-9])__(?P<bold2>\S(?:.*?\S)?)__(?![A-Za-z0-9])"
    r"|~~(?P<strike>\S(?:.*?\S)?)~~"
    r"|\*(?P<italic>[^\s*](?:.*?[^\s*])?)\*"
    r"|(?<![A-Za-z0-9_])_(?P<italic2>\S(?:.*?\S)?)_(?![A-Za-z0-9_])",
    re.DOTALL,
)


def parse_inline(text: str, style: Style | None = None) -> list[Span]:
    """Turn inline Markdown into styled spans, recursing through nesting."""
    style = style or Style()
    spans: list[Span] = []
    position = 0

    def literal(chunk: str) -> None:
        for piece_index, piece in enumerate(chunk.split("\n")):
            if piece_index:
                spans.append(Span(kind="break", style=style))
            if piece:
                spans.append(Span(text=piece, style=style))

    for match in _INLINE_RE.finditer(text):
        if match.start() > position:
            literal(text[position : match.start()])
        position = match.end()

        if match.group("escape"):
            literal(match.group("escape")[1])
        elif match.group("code") is not None:
            inner = match.group("code").strip()
            spans.append(Span(text=inner, style=_replace(style, code=True)))
        elif match.group("src") is not None:
            spans.append(
                Span(
                    kind="image",
                    text=_unescape(match.group("alt")),
                    src=match.group("src"),
                    style=style,
                )
            )
        elif match.group("href") is not None:
            href = _link_target(match.group("href"))
            label = match.group("label") or href
            spans.extend(parse_inline(label, _replace(style, href=href)))
        elif match.group("auto"):
            url = match.group("auto")
            spans.append(Span(text=url, style=_replace(style, href=url)))
        elif match.group("hardbreak"):
            spans.append(Span(kind="break", style=style))
        elif match.group("vtag"):
            vertical = "superscript" if match.group("vtag") == "sup" else "subscript"
            spans.extend(
                parse_inline(match.group("vtext"), _replace(style, vertical=vertical))
            )
        elif match.group("html") is not None:
            continue  # inline HTML Word has no equivalent for
        elif match.group("bi") is not None:
            spans.extend(
                parse_inline(match.group("bi"), _replace(style, bold=True, italic=True))
            )
        elif match.group("bold") is not None or match.group("bold2") is not None:
            inner = match.group("bold")
            inner = match.group("bold2") if inner is None else inner
            spans.extend(parse_inline(inner, _replace(style, bold=True)))
        elif match.group("strike") is not None:
            spans.extend(parse_inline(match.group("strike"), _replace(style, strike=True)))
        else:
            inner = match.group("italic")
            inner = match.group("italic2") if inner is None else inner
            spans.extend(parse_inline(inner, _replace(style, italic=True)))

    if position < len(text):
        literal(text[position:])
    return spans


def _link_target(destination: str) -> str:
    """Drop an optional link title and the angle brackets around a target."""
    destination = destination.strip()
    for quote_char in ('"', "'"):
        marker = f" {quote_char}"
        if marker in destination and destination.endswith(quote_char):
            destination = destination[: destination.index(marker)]
            break
    if destination.startswith("<") and destination.endswith(">"):
        destination = destination[1:-1]
    return destination.strip()


def _unescape(text: str) -> str:
    return re.sub(r"\\([\\`*_{}\[\]()#+\-.!|~<>])", r"\1", text)


# ------------------------------------------------------------------- writer


class _Writer:
    """Render parsed blocks into a python-docx document."""

    def __init__(
        self,
        base_dir: Path,
        embed_images: bool,
        warnings: list[str],
        settings: DocxSettings | None = None,
    ):
        self.document = Document()
        self.base_dir = base_dir
        self.embed_images = embed_images
        self.warnings = warnings
        self.settings = settings or DocxSettings()
        numbering_part = self.document.part.numbering_part
        self._numbering = numbering_part.numbering_definitions._numbering
        self._page_is_empty = True
        self._prepare_page()
        self._prepare_fonts()
        self._prepare_styles()

    # ------------------------------------------------------------- setup

    def _prepare_page(self) -> None:
        section = self.document.sections[0]
        section.page_width, section.page_height = map(Emu, self.settings.page)
        self.usable_width = (
            section.page_width - section.left_margin - section.right_margin
        )

    def _prepare_fonts(self) -> None:
        """Put the whole document on one font.

        Setting `Normal` is not enough: the built-in headings resolve their
        font through the theme, so they would keep coming out in the template's
        Calibri while the body text obeyed the chosen face.
        """
        font = self.settings.body_font
        self._retheme_fonts(font)
        normal = self.document.styles["Normal"]
        normal.font.size = self.settings.body_size
        normal.paragraph_format.line_spacing = self.settings.line_spacing
        _name_font(normal, font)
        _paint_black(normal)
        for name in _THEMED_STYLES:
            try:
                style = self.document.styles[name]
            except KeyError:  # a template without that built-in style
                continue
            _name_font(style, font)
            _paint_black(style)

    def _retheme_fonts(self, name: str) -> None:
        """Rename the theme's major/minor faces, which most styles point at."""
        part = next(
            (
                candidate
                for candidate in self.document.part.package.iter_parts()
                if str(candidate.partname).endswith("theme1.xml")
            ),
            None,
        )
        if part is None:
            return
        try:
            theme = part.blob.decode("utf-8")
        except Exception:  # noqa: BLE001 - a theme we cannot read is not fatal
            return
        part._blob = _THEME_FACE_RE.sub(
            lambda match: match.group(1) + name, theme
        ).encode("utf-8")

    def _prepare_styles(self) -> None:
        styles = self.document.styles
        names = {style.name for style in styles}

        code_size = Pt(round(self.settings.font_size * _CODE_SIZE_RATIO, 1))

        if "Code" not in names:
            code = styles.add_style("Code", WD_STYLE_TYPE.PARAGRAPH)
            code.base_style = styles["Normal"]
            code.font.size = code_size
            _name_font(code, _CODE_FONT)
            code.paragraph_format.space_before = Pt(0)
            code.paragraph_format.space_after = Pt(0)
            code.paragraph_format.left_indent = Pt(12)
            code.paragraph_format.line_spacing = 1.0
            shading = f'<w:shd {nsdecls("w")} w:val="clear" w:color="auto" w:fill="F4F4F4"/>'
            code.element.get_or_add_pPr().append(parse_xml(shading))

        if "Code Char" not in names:
            inline_code = styles.add_style("Code Char", WD_STYLE_TYPE.CHARACTER)
            inline_code.font.size = code_size
            _name_font(inline_code, _CODE_FONT)

        if "Hyperlink" not in names:
            # Black like the rest of the text; the underline still marks it out.
            link = styles.add_style("Hyperlink", WD_STYLE_TYPE.CHARACTER)
            link.font.color.rgb = RGBColor(0x00, 0x00, 0x00)
            link.font.underline = True

    # -------------------------------------------------------------- blocks

    def write(self, blocks: list, level: int = 0, indent: int = 0) -> None:
        for block in blocks:
            self._write_block(block, level, indent)
            self._page_is_empty = False

    def add_title(self, text: str) -> None:
        paragraph = self.document.add_paragraph(style="Title")
        self._add_spans(paragraph, parse_inline(text))
        self._page_is_empty = False

    def add_table_of_contents(self) -> None:
        """Drop in Word's own TOC field and ask Word to fill it on open.

        The field is empty until it updates, so `updateFields` matters: without
        it the reader sees a blank page where the contents should be.
        """
        heading = self.document.add_paragraph(style="Heading 1")
        heading.add_run("Mục lục")

        paragraph = self.document.add_paragraph(style="Normal")
        _add_field_char(paragraph, "begin")
        instruction = paragraph.add_run()._r
        instruction.append(
            parse_xml(
                f'<w:instrText {nsdecls("w")} xml:space="preserve">'
                r' TOC \o "1-3" \h \z \u '
                "</w:instrText>"
            )
        )
        _add_field_char(paragraph, "separate")
        paragraph.add_run("Nhấn F9 trong Word để cập nhật mục lục.")
        _add_field_char(paragraph, "end")

        self.document.add_paragraph().add_run().add_break(WD_BREAK.PAGE)

        settings = self.document.settings.element
        if settings.find(qn("w:updateFields")) is None:
            settings.append(parse_xml(f'<w:updateFields {nsdecls("w")} w:val="true"/>'))
        self._page_is_empty = True  # the contents end on a page of their own

    def _write_block(self, block, level: int, indent: int) -> None:
        if isinstance(block, Heading):
            paragraph = self._paragraph(f"Heading {min(block.level, 6)}", indent)
            # A break at the top of an empty page would only add a blank one.
            if (
                self.settings.page_break_before_h1
                and block.level == 1
                and not self._page_is_empty
            ):
                paragraph.paragraph_format.page_break_before = True
            self._add_spans(paragraph, parse_inline(block.text))
        elif isinstance(block, Paragraph):
            paragraph = self._paragraph("Normal", indent)
            self._add_spans(paragraph, parse_inline(block.text))
        elif isinstance(block, CodeBlock):
            self._write_code(block, indent)
        elif isinstance(block, Quote):
            self._write_quote(block, level, indent)
        elif isinstance(block, ListBlock):
            self._write_list(block, level, indent)
        elif isinstance(block, Table):
            self._write_table(block)
        elif isinstance(block, Rule):
            self._write_rule(indent)

    def _paragraph(self, style: str, indent: int = 0):
        paragraph = self.document.add_paragraph(style=style)
        if indent:
            paragraph.paragraph_format.left_indent = Emu(indent * _TWIP)
        return paragraph

    def _write_code(self, block: CodeBlock, indent: int) -> None:
        # One paragraph per line: DOCX_STYLE_MAP folds adjacent `Code`
        # paragraphs back into a single fenced block on the way out.
        for line in (block.text or "").split("\n"):
            paragraph = self._paragraph("Code", indent)
            if line:
                paragraph.add_run(_clean(line))

    def _write_quote(self, block: Quote, level: int, indent: int) -> None:
        inner = indent + _HANGING_INDENT
        for child in block.blocks:
            if isinstance(child, Paragraph):
                paragraph = self._paragraph("Quote", inner)
                self._add_spans(paragraph, parse_inline(child.text))
            else:
                self._write_block(child, level, inner)

    def _write_list(self, block: ListBlock, level: int, indent: int) -> None:
        num_id = None
        depth = min(level, 8)
        child_indent = indent + _LEVEL_INDENT * (depth + 1)

        for item in block.items:
            blocks = item.blocks or [Paragraph("")]
            first, *rest = blocks

            # A `- - text` item carries no text of its own: Word lists that
            # start below level one arrive as an outer item wrapping nothing
            # but a nested list. Bulleting it would leave a stray, empty
            # bullet above the nested items, so hoist the nested list into
            # this level instead - the content takes the bullet the empty
            # item would have had.
            hoisted = isinstance(first, ListBlock) and item.checked is None
            if hoisted:
                rest = blocks
            else:
                if num_id is None:
                    num_id = self._new_numbering(block.ordered, block.start)
                paragraph = self._paragraph("List Paragraph")
                self._apply_numbering(paragraph, num_id, depth, indent)
                if item.checked is not None:
                    paragraph.add_run("☒ " if item.checked else "☐ ")
                if isinstance(first, Paragraph):
                    self._add_spans(paragraph, parse_inline(first.text))
                else:
                    self._write_block(first, level + 1, child_indent)

            inner_level = level if hoisted else level + 1
            for child in rest:
                if isinstance(child, ListBlock):
                    self._write_list(child, inner_level, indent)
                else:
                    self._write_block(child, inner_level, child_indent)

    def _write_table(self, block: Table) -> None:
        width = len(block.header)
        if not width:
            return
        table = self.document.add_table(rows=1, cols=width)
        table.style = self.document.styles["Table Grid"]
        table.alignment = WD_TABLE_ALIGNMENT.LEFT

        header_cells = table.rows[0].cells
        for column, text in enumerate(block.header):
            self._fill_cell(header_cells[column], text, block.alignments[column], True)
        self._repeat_header(table.rows[0])

        for row in block.rows:
            cells = table.add_row().cells
            for column, text in enumerate(row[:width]):
                self._fill_cell(cells[column], text, block.alignments[column])

    def _fill_cell(self, cell, text: str, alignment: str, bold: bool = False) -> None:
        paragraph = cell.paragraphs[0]
        if alignment == "center":
            paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        elif alignment == "right":
            paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        spans = parse_inline(text)
        if bold:
            spans = [
                _replace(span, style=_replace(span.style, bold=True)) for span in spans
            ]
        self._add_spans(paragraph, spans)

    def _repeat_header(self, row) -> None:
        properties = row._tr.get_or_add_trPr()
        properties.append(parse_xml(f'<w:tblHeader {nsdecls("w")} w:val="true"/>'))

    def _write_rule(self, indent: int) -> None:
        paragraph = self._paragraph("Normal", indent)
        paragraph._p.get_or_add_pPr().append(
            parse_xml(
                f'<w:pBdr {nsdecls("w")}>'
                '<w:bottom w:val="single" w:sz="6" w:space="1" w:color="A0A0A0"/>'
                "</w:pBdr>"
            )
        )

    # ---------------------------------------------------------- numbering

    def _new_numbering(self, ordered: bool, start: int) -> int:
        """Give every list its own numbering instance so counters restart."""
        defined = self._numbering.findall(qn("w:abstractNum"))
        abstract_id = max(
            (int(node.get(qn("w:abstractNumId"))) for node in defined), default=-1
        ) + 1
        levels = "".join(self._level_xml(level, ordered, start) for level in range(9))
        abstract = parse_xml(
            f'<w:abstractNum {nsdecls("w")} w:abstractNumId="{abstract_id}">'
            f"{levels}</w:abstractNum>"
        )
        if defined:
            defined[-1].addnext(abstract)
        else:
            self._numbering.insert(0, abstract)
        return self._numbering.add_num(abstract_id).numId

    def _level_xml(self, level: int, ordered: bool, start: int) -> str:
        left = _LEVEL_INDENT * (level + 1)
        if ordered:
            fmt = _ORDERED_FORMATS[level % len(_ORDERED_FORMATS)]
            text = f"%{level + 1}."
            fonts = ""
        else:
            fmt = "bullet"
            text = _BULLET_CHARS[level % len(_BULLET_CHARS)]
            font = _BULLET_FONTS[level % len(_BULLET_FONTS)]
            fonts = (
                f'<w:rPr><w:rFonts w:ascii="{font}" w:hAnsi="{font}" '
                'w:hint="default"/></w:rPr>'
            )
        return (
            f'<w:lvl w:ilvl="{level}">'
            f'<w:start w:val="{start if level == 0 else 1}"/>'
            f'<w:numFmt w:val="{fmt}"/>'
            f'<w:lvlText w:val="{text}"/>'
            '<w:lvlJc w:val="left"/>'
            f'<w:pPr><w:ind w:left="{left}" w:hanging="{_HANGING_INDENT}"/></w:pPr>'
            f"{fonts}</w:lvl>"
        )

    def _apply_numbering(self, paragraph, num_id: int, level: int, indent: int) -> None:
        properties = paragraph._p.get_or_add_pPr()
        properties.append(
            parse_xml(
                f'<w:numPr {nsdecls("w")}>'
                f'<w:ilvl w:val="{level}"/><w:numId w:val="{num_id}"/></w:numPr>'
            )
        )
        if indent:
            left = indent + _LEVEL_INDENT * (level + 1)
            properties.append(
                parse_xml(
                    f'<w:ind {nsdecls("w")} w:left="{left}" '
                    f'w:hanging="{_HANGING_INDENT}"/>'
                )
            )

    # -------------------------------------------------------------- inline

    def _add_spans(self, paragraph, spans: list[Span]) -> None:
        for href, group in groupby(spans, key=lambda span: span.style.href):
            container = self._hyperlink(paragraph, href) if href else None
            for span in group:
                run = self._add_span(paragraph, span)
                if container is not None and run is not None:
                    container.append(run._r)

    def _add_span(self, paragraph, span: Span):
        if span.kind == "break":
            run = paragraph.add_run()
            run.add_break()
            return run
        if span.kind == "image":
            return self._add_image(paragraph, span)

        run = paragraph.add_run(_clean(span.text))
        style = span.style
        if style.bold:
            run.bold = True
        if style.italic:
            run.italic = True
        if style.strike:
            run.font.strike = True
        if style.vertical == "superscript":
            run.font.superscript = True
        elif style.vertical == "subscript":
            run.font.subscript = True
        if style.code:
            run.style = self.document.styles["Code Char"]
        elif style.href:
            run.style = self.document.styles["Hyperlink"]
        return run

    def _add_image(self, paragraph, span: Span):
        source = self._resolve(span.src) if self.embed_images else None
        if source is None:
            label = span.text or span.src
            return paragraph.add_run(_clean(label)) if label else None
        try:
            run = paragraph.add_run()
            picture = run.add_picture(str(source))
        except Exception as exc:  # noqa: BLE001 - one bad image must not kill the file
            self.warnings.append(f"Không chèn được ảnh {span.src}: {exc}")
            return paragraph.add_run(_clean(span.text or span.src))
        if picture.width > self.usable_width:
            ratio = self.usable_width / picture.width
            picture.width = Emu(int(picture.width * ratio))
            picture.height = Emu(int(picture.height * ratio))
        return run

    def _resolve(self, src: str) -> Path | None:
        if not src:
            return None
        if urlparse(src).scheme in ("http", "https", "data", "ftp"):
            self.warnings.append(f"Bỏ qua ảnh ở xa (không tải về): {src}")
            return None
        candidate = Path(unquote(src))
        if not candidate.is_absolute():
            candidate = self.base_dir / candidate
        if candidate.is_file():
            return candidate
        self.warnings.append(f"Không tìm thấy ảnh: {src}")
        return None

    def _hyperlink(self, paragraph, url: str):
        element = OxmlElement("w:hyperlink")
        try:
            relationship = paragraph.part.relate_to(url, RT.HYPERLINK, is_external=True)
        except Exception:  # noqa: BLE001 - an unusable target still keeps its text
            return None
        element.set(qn("r:id"), relationship)
        paragraph._p.append(element)
        return element


def _add_field_char(paragraph, kind: str) -> None:
    paragraph.add_run()._r.append(
        parse_xml(f'<w:fldChar {nsdecls("w")} w:fldCharType="{kind}"/>')
    )


def _name_font(style, name: str) -> None:
    _name_font_element(style.element, name)


def _name_font_element(element, name: str) -> None:
    """Name the font for every script, and drop the theme reference.

    `Font.name` only fills in the Latin face; a theme reference left next to it
    wins in Word, which is how headings keep the template's font.
    """
    fonts = element.get_or_add_rPr().get_or_add_rFonts()
    for attribute in ("asciiTheme", "hAnsiTheme", "eastAsiaTheme", "cstheme"):
        fonts.attrib.pop(qn(f"w:{attribute}"), None)
    for attribute in ("ascii", "hAnsi", "eastAsia", "cs"):
        fonts.set(qn(f"w:{attribute}"), name)


def _paint_black(style) -> None:
    """Force plain black text, theme colour and all.

    Word's built-in headings are blue by way of `themeColor="accent1"`, which
    outranks a colour set next to it, so the reference has to go. `Title` and
    `Intense Quote` also draw an accent-coloured rule, which would otherwise be
    the one blue thing left on the page.
    """
    color = style.element.get_or_add_rPr().get_or_add_color()
    _drop_theme_color(color)
    color.set(qn("w:val"), "000000")

    properties = style.element.find(qn("w:pPr"))
    borders = None if properties is None else properties.find(qn("w:pBdr"))
    for edge in borders if borders is not None else ():
        _drop_theme_color(edge)
        edge.set(qn("w:color"), "000000")


def _drop_theme_color(element) -> None:
    for attribute in ("themeColor", "themeTint", "themeShade"):
        element.attrib.pop(qn(f"w:{attribute}"), None)


def _clean(text: str) -> str:
    return _CONTROL_CHARS.sub("", text)


# --------------------------------------------------------------- public api


def markdown_text_to_docx(
    text: str,
    destination: str | Path,
    base_dir: str | Path | None = None,
    embed_images: bool = True,
    title: str | None = None,
    settings: DocxSettings | None = None,
) -> list[str]:
    """Write `text` as a .docx and return the warnings collected on the way."""
    destination = Path(destination)
    warnings: list[str] = []
    settings = settings or DocxSettings()
    blocks, metadata = parse_markdown(text)

    writer = _Writer(
        Path(base_dir or destination.parent), embed_images, warnings, settings
    )
    _apply_metadata(writer.document, metadata)

    # The document only gets a title of its own when the Markdown has no H1 to
    # act as one - mirroring the H1 the forward conversion adds.
    heading = metadata.get("title") or title
    starts_with_h1 = bool(blocks) and isinstance(blocks[0], Heading) and blocks[0].level == 1
    if heading and not starts_with_h1:
        writer.add_title(heading)

    if settings.table_of_contents:
        writer.add_table_of_contents()

    writer.write(blocks)
    destination.parent.mkdir(parents=True, exist_ok=True)
    writer.document.save(str(destination))
    return warnings


def _apply_metadata(document, metadata: dict[str, str]) -> None:
    properties = document.core_properties
    for key, attribute in (
        ("title", "title"),
        ("author", "author"),
        ("subject", "subject"),
        ("keywords", "keywords"),
        ("description", "comments"),
        ("category", "category"),
    ):
        value = metadata.get(key)
        if value:
            setattr(properties, attribute, value)


def markdown_to_docx(
    source: str | Path,
    destination: str | Path,
    embed_images: bool = True,
    add_title_heading: bool = True,
    settings: DocxSettings | None = None,
) -> list[str]:
    """Convert one .md file into `destination` (.docx)."""
    source = Path(source)
    text = source.read_text(encoding="utf-8-sig", errors="replace")
    return markdown_text_to_docx(
        text,
        destination,
        base_dir=source.parent,
        embed_images=embed_images,
        title=source.stem if add_title_heading else None,
        settings=settings,
    )


__all__ = [
    "DEFAULT_FONT",
    "DEFAULT_FONT_SIZE",
    "MARKDOWN_EXTENSIONS",
    "PAGE_SIZES",
    "DocxSettings",
    "markdown_text_to_docx",
    "markdown_to_docx",
    "parse_inline",
    "parse_markdown",
]
