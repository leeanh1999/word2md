"""Tests for the Markdown -> Word direction."""

from __future__ import annotations

import base64
import sys
import tempfile
import unittest
from pathlib import Path

sys.path.insert(0, str(Path(__file__).resolve().parents[1]))

from src.converter import (  # noqa: E402
    STATUS_SUCCESS,
    ConversionOptions,
    collect_files,
    convert_file,
    docx_to_markdown,
    output_suffix,
    target_path,
)
from src.md_to_docx import (  # noqa: E402
    DEFAULT_FONT,
    DEFAULT_FONT_SIZE,
    PAGE_SIZES,
    CodeBlock,
    DocxSettings,
    ListBlock,
    Quote,
    Rule,
    Table,
    markdown_text_to_docx,
    markdown_to_docx,
    parse_inline,
    parse_markdown,
)

# A 10x10 PNG, small enough to keep in the test file.
PNG = base64.b64decode(
    "iVBORw0KGgoAAAANSUhEUgAAAAoAAAAKCAYAAACNMs+9AAAAHUlEQVR42mP8z8BQz0AEYBxVSF+F"
    "/6EKRxUOKgUA7WMH8V0F0kIAAAAASUVORK5CYII="
)


def blocks_of(text: str):
    return parse_markdown(text)[0]


def _font_of(style) -> str:
    """The Latin face a style names outright, theme references aside."""
    from docx.oxml.ns import qn

    fonts = style.element.get_or_add_rPr().get_or_add_rFonts()
    return fonts.get(qn("w:ascii")) or ""


class BlockParserTests(unittest.TestCase):
    def test_headings(self):
        blocks = blocks_of("# Một\n\n### Ba ###\n\nSetext\n======\n")
        self.assertEqual([(b.level, b.text) for b in blocks], [(1, "Một"), (3, "Ba"), (1, "Setext")])

    def test_paragraph_joins_wrapped_lines(self):
        (block,) = blocks_of("dòng một\ndòng hai\n")
        self.assertEqual(block.text, "dòng một dòng hai")

    def test_hard_break_survives_as_newline(self):
        (block,) = blocks_of("dòng một  \ndòng hai\n")
        self.assertEqual(block.text, "dòng một\ndòng hai")

    def test_nested_list(self):
        (block,) = blocks_of("- một\n- hai\n  - hai.a\n  - hai.b\n- ba\n")
        self.assertIsInstance(block, ListBlock)
        self.assertFalse(block.ordered)
        self.assertEqual(len(block.items), 3)
        nested = block.items[1].blocks[1]
        self.assertIsInstance(nested, ListBlock)
        self.assertEqual([item.blocks[0].text for item in nested.items], ["hai.a", "hai.b"])

    def test_ordered_list_keeps_its_start(self):
        (block,) = blocks_of("3. ba\n4. bốn\n")
        self.assertTrue(block.ordered)
        self.assertEqual(block.start, 3)

    def test_loose_list_stays_one_list(self):
        blocks = blocks_of("1. một\n2. hai\n\n   Ghi chú.\n\n3. ba\n")
        self.assertEqual(len(blocks), 1)
        self.assertEqual(len(blocks[0].items), 3)
        self.assertEqual(blocks[0].items[1].blocks[1].text, "Ghi chú.")

    def test_task_list(self):
        (block,) = blocks_of("- [ ] chưa\n- [x] rồi\n")
        self.assertEqual([item.checked for item in block.items], [False, True])

    def test_table(self):
        (block,) = blocks_of("| A | B | C |\n| --- | :-: | --: |\n| 1 | 2 | 3 |\n")
        self.assertIsInstance(block, Table)
        self.assertEqual(block.header, ["A", "B", "C"])
        self.assertEqual(block.rows, [["1", "2", "3"]])
        self.assertEqual(block.alignments, ["", "center", "right"])

    def test_table_cell_keeps_escaped_pipe(self):
        (block,) = blocks_of("| A | B |\n| --- | --- |\n| a \\| b | c |\n")
        self.assertEqual(block.rows[0], ["a | b", "c"])

    def test_ragged_table_is_padded(self):
        (block,) = blocks_of("| A | B | C |\n| --- | --- | --- |\n| 1 |\n")
        self.assertEqual(block.rows[0], ["1", "", ""])

    def test_fenced_code_keeps_indentation(self):
        (block,) = blocks_of("```python\ndef f():\n    return 1\n```\n")
        self.assertIsInstance(block, CodeBlock)
        self.assertEqual(block.language, "python")
        self.assertEqual(block.text, "def f():\n    return 1")

    def test_quote_and_rule(self):
        blocks = blocks_of("> trích\n> dẫn\n\n---\n")
        self.assertIsInstance(blocks[0], Quote)
        self.assertEqual(blocks[0].blocks[0].text, "trích dẫn")
        self.assertIsInstance(blocks[1], Rule)

    def test_rule_is_not_a_list(self):
        self.assertIsInstance(blocks_of("* * *\n")[0], Rule)

    def test_front_matter_is_metadata_not_content(self):
        blocks, metadata = parse_markdown("---\ntitle: Tiêu đề\nauthor: Anh\n---\n\nNội dung.\n")
        self.assertEqual(metadata, {"title": "Tiêu đề", "author": "Anh"})
        self.assertEqual([b.text for b in blocks], ["Nội dung."])

    def test_empty_document(self):
        self.assertEqual(blocks_of("\n\n"), [])


class InlineParserTests(unittest.TestCase):
    def styles(self, text):
        return [(span.text, span.style) for span in parse_inline(text)]

    def test_emphasis(self):
        spans = parse_inline("a **đậm** b *nghiêng* c ***cả hai***")
        bold = [s for s in spans if s.style.bold and not s.style.italic]
        italic = [s for s in spans if s.style.italic and not s.style.bold]
        both = [s for s in spans if s.style.bold and s.style.italic]
        self.assertEqual([s.text for s in bold], ["đậm"])
        self.assertEqual([s.text for s in italic], ["nghiêng"])
        self.assertEqual([s.text for s in both], ["cả hai"])

    def test_underscore_inside_a_word_is_literal(self):
        spans = parse_inline("snake_case_name")
        self.assertEqual([s.text for s in spans], ["snake_case_name"])
        self.assertFalse(any(s.style.italic for s in spans))

    def test_code_and_strike(self):
        spans = parse_inline("`mã` và ~~bỏ~~")
        self.assertTrue(next(s for s in spans if s.text == "mã").style.code)
        self.assertTrue(next(s for s in spans if s.text == "bỏ").style.strike)

    def test_link_and_autolink(self):
        (span,) = [s for s in parse_inline("[nhãn](https://x.dev/a%20b)") if s.text]
        self.assertEqual(span.text, "nhãn")
        self.assertEqual(span.style.href, "https://x.dev/a%20b")
        (auto,) = [s for s in parse_inline("<https://x.dev>") if s.text]
        self.assertEqual(auto.style.href, "https://x.dev")

    def test_link_title_is_dropped(self):
        (span,) = [s for s in parse_inline('[a](x.md "chú thích")') if s.text]
        self.assertEqual(span.style.href, "x.md")

    def test_image_span(self):
        (span,) = parse_inline("![sơ đồ](img/a.png)")
        self.assertEqual((span.kind, span.text, span.src), ("image", "sơ đồ", "img/a.png"))

    def test_bold_inside_link_keeps_both(self):
        spans = [s for s in parse_inline("[**đậm**](x.md)") if s.text]
        self.assertEqual(len(spans), 1)
        self.assertTrue(spans[0].style.bold)
        self.assertEqual(spans[0].style.href, "x.md")

    def test_superscript_and_subscript(self):
        spans = parse_inline("H<sub>2</sub>O<sup>x</sup>")
        self.assertEqual(
            [(s.text, s.style.vertical) for s in spans],
            [("H", ""), ("2", "subscript"), ("O", ""), ("x", "superscript")],
        )

    def test_escaped_markers_stay_literal(self):
        spans = parse_inline(r"2 \* 3 \*\* 4")
        self.assertEqual("".join(s.text for s in spans), "2 * 3 ** 4")

    def test_hard_break_becomes_a_break_span(self):
        kinds = [s.kind for s in parse_inline("a\nb")]
        self.assertEqual(kinds, ["text", "break", "text"])


class RoundTripTests(unittest.TestCase):
    """Markdown -> .docx -> Markdown must come back recognisable."""

    @classmethod
    def setUpClass(cls):
        cls.tmp = tempfile.TemporaryDirectory()
        cls.root = Path(cls.tmp.name)

    @classmethod
    def tearDownClass(cls):
        cls.tmp.cleanup()

    def round_trip(self, markdown: str, name: str = "rt") -> str:
        destination = self.root / f"{name}.docx"
        markdown_text_to_docx(markdown, destination, base_dir=self.root)
        text, _ = docx_to_markdown(
            destination,
            ConversionOptions(extract_images=False, add_title_heading=False),
        )
        return text

    def test_structure_survives(self):
        source = (
            "# Tiêu đề\n\n"
            "Đoạn **đậm** *nghiêng* `mã` [liên kết](https://x.dev).\n\n"
            "## Danh sách\n\n"
            "- một\n- hai\n  - hai.a\n\n"
            "1. đầu\n2. cuối\n\n"
            "> trích dẫn\n\n"
            "```\nprint(1)\n```\n"
        )
        text = self.round_trip(source, "structure")
        for expected in (
            "# Tiêu đề",
            "## Danh sách",
            "**đậm**",
            "*nghiêng*",
            "`mã`",
            "[liên kết](https://x.dev)",
            "- một",
            "  - hai.a",
            "1. đầu",
            "> trích dẫn",
            "```\nprint(1)\n```",
        ):
            self.assertIn(expected, text)

    def test_table_survives(self):
        text = self.round_trip(
            "| Tên | Số |\n| --- | --- |\n| a | 1 |\n| b | 2 |\n", "table"
        )
        self.assertIn("| --- | --- |", text)
        self.assertIn("| a | 1 |", text)
        self.assertIn("| b | 2 |", text)

    def test_vietnamese_text_is_intact(self):
        text = self.round_trip("Chào bạn, đây là tiếng Việt có dấu.\n", "vn")
        self.assertIn("Chào bạn, đây là tiếng Việt có dấu.", text)

    def test_image_is_embedded_and_comes_back(self):
        folder = self.root / "with_images"
        (folder / "doc_images").mkdir(parents=True, exist_ok=True)
        (folder / "doc_images" / "image1.png").write_bytes(PNG)
        source = folder / "doc.md"
        source.write_text("# Ảnh\n\n![sơ đồ](doc_images/image1.png)\n", encoding="utf-8")

        destination = folder / "doc.docx"
        self.assertEqual(markdown_to_docx(source, destination), [])
        text, _ = docx_to_markdown(
            destination,
            ConversionOptions(add_title_heading=False),
            image_dir=folder / "back_images",
        )
        self.assertIn("![](back_images/image1.png)", text)
        self.assertTrue((folder / "back_images" / "image1.png").is_file())

    def test_missing_image_warns_and_keeps_the_alt_text(self):
        destination = self.root / "missing.docx"
        warnings = markdown_text_to_docx(
            "![thiếu](nowhere/x.png)\n", destination, base_dir=self.root
        )
        self.assertTrue(any("Không tìm thấy ảnh" in w for w in warnings))
        text, _ = docx_to_markdown(
            destination, ConversionOptions(add_title_heading=False)
        )
        self.assertIn("thiếu", text)

    def test_remote_image_is_reported_not_fetched(self):
        warnings = markdown_text_to_docx(
            "![xa](https://example.com/a.png)\n",
            self.root / "remote.docx",
            base_dir=self.root,
        )
        self.assertTrue(any("ảnh ở xa" in w for w in warnings))

    def test_front_matter_fills_the_document_properties(self):
        destination = self.root / "meta.docx"
        markdown_text_to_docx(
            "---\ntitle: Báo cáo\nauthor: Anh\n---\n\nNội dung.\n",
            destination,
            base_dir=self.root,
        )
        import docx

        properties = docx.Document(str(destination)).core_properties
        self.assertEqual(properties.title, "Báo cáo")
        self.assertEqual(properties.author, "Anh")


class SettingsTests(unittest.TestCase):
    @classmethod
    def setUpClass(cls):
        cls.tmp = tempfile.TemporaryDirectory()
        cls.root = Path(cls.tmp.name)

    @classmethod
    def tearDownClass(cls):
        cls.tmp.cleanup()

    def build(self, markdown: str, name: str, **kwargs):
        import docx

        destination = self.root / f"{name}.docx"
        markdown_text_to_docx(
            markdown, destination, base_dir=self.root, settings=DocxSettings(**kwargs)
        )
        return docx.Document(str(destination))

    def test_times_new_roman_is_the_default(self):
        document = self.build("# Tiêu đề\n\nNội dung.\n", "default")
        for name in ("Normal", "Heading 1", "Title"):
            self.assertEqual(
                _font_of(document.styles[name]),
                DEFAULT_FONT,
                f"style {name} kept the template font",
            )
        self.assertEqual(document.styles["Normal"].font.size.pt, DEFAULT_FONT_SIZE)

    def test_every_style_writes_in_black(self):
        from docx.oxml.ns import qn

        document = self.build("# Tiêu đề\n\n> trích dẫn\n", "black")
        for name in (
            "Normal",
            "Title",
            "Quote",
            "Intense Quote",
            "Hyperlink",
            *(f"Heading {level}" for level in range(1, 7)),
        ):
            color = document.styles[name].element.get_or_add_rPr().get_or_add_color()
            self.assertEqual(color.get(qn("w:val")), "000000", name)
            self.assertIsNone(color.get(qn("w:themeColor")), f"{name} still themed")

        # Title's accent rule would be the one blue thing left on the page.
        self.assertNotIn("4F81BD", document.styles["Title"].element.xml)

    def test_an_empty_font_falls_back_to_the_default(self):
        self.assertEqual(DocxSettings(font="   ").body_font, DEFAULT_FONT)
        document = self.build("Nội dung.\n", "blank_font", font="")
        self.assertEqual(_font_of(document.styles["Normal"]), DEFAULT_FONT)

    def test_chosen_font_size_and_spacing(self):
        document = self.build(
            "Nội dung.\n", "chosen", font="Arial", font_size=14, line_spacing=1.5
        )
        normal = document.styles["Normal"]
        self.assertEqual(_font_of(normal), "Arial")
        self.assertEqual(normal.font.size.pt, 14)
        self.assertEqual(normal.paragraph_format.line_spacing, 1.5)
        # Code keeps its monospace face whatever the body font is.
        self.assertEqual(_font_of(document.styles["Code"]), "Consolas")

    def test_page_size(self):
        letter = self.build("x\n", "letter", page_size="Letter").sections[0]
        self.assertEqual(
            (letter.page_width, letter.page_height), PAGE_SIZES["Letter"]
        )
        a4 = self.build("x\n", "a4", page_size="Nonsense").sections[0]
        self.assertEqual((a4.page_width, a4.page_height), PAGE_SIZES["A4"])

    def test_page_break_before_h1_skips_the_first_one(self):
        document = self.build(
            "# Một\n\na\n\n# Hai\n\nb\n", "breaks", page_break_before_h1=True
        )
        breaks = [
            paragraph.paragraph_format.page_break_before
            for paragraph in document.paragraphs
            if paragraph.style.name == "Heading 1"
        ]
        self.assertEqual(breaks, [None, True])

    def test_table_of_contents_field(self):
        import docx
        from docx.oxml.ns import qn

        destination = self.root / "toc.docx"
        markdown_text_to_docx(
            "# Một\n\na\n",
            destination,
            base_dir=self.root,
            settings=DocxSettings(table_of_contents=True),
        )
        document = docx.Document(str(destination))
        self.assertIn("Mục lục", [p.text for p in document.paragraphs])
        self.assertIn('TOC \\o "1-3"', document.element.xml)
        self.assertIsNotNone(document.settings.element.find(qn("w:updateFields")))


class DispatcherTests(unittest.TestCase):
    @classmethod
    def setUpClass(cls):
        cls.tmp = tempfile.TemporaryDirectory()
        cls.root = Path(cls.tmp.name)
        cls.source = cls.root / "ghi chú.md"
        cls.source.write_text("# Ghi chú\n\nNội dung.\n", encoding="utf-8")

    @classmethod
    def tearDownClass(cls):
        cls.tmp.cleanup()

    def test_output_suffix(self):
        self.assertEqual(output_suffix(Path("a.md")), ".docx")
        self.assertEqual(output_suffix(Path("a.MARKDOWN")), ".docx")
        self.assertEqual(output_suffix(Path("a.docx")), ".md")
        self.assertEqual(output_suffix(Path("a.xlsx")), ".md")

    def test_target_path_follows_the_direction(self):
        folder = self.root / "targets"
        self.assertEqual(target_path(Path("a.md"), folder).suffix, ".docx")
        self.assertEqual(target_path(Path("a.docx"), folder).suffix, ".md")

    def test_convert_file_writes_a_docx(self):
        folder = self.root / "out"
        result = convert_file(self.source, folder)
        self.assertEqual(result.status, STATUS_SUCCESS, result.message)
        self.assertEqual(result.output.name, "ghi chú.docx")
        text, _ = docx_to_markdown(
            result.output, ConversionOptions(add_title_heading=False)
        )
        self.assertIn("# Ghi chú", text)
        self.assertIn("Nội dung.", text)

    def test_no_overwrite_creates_unique_names(self):
        folder = self.root / "unique"
        first = convert_file(self.source, folder)
        second = convert_file(self.source, folder)
        self.assertNotEqual(first.output, second.output)
        self.assertTrue(second.output.name.endswith("(1).docx"))

    def test_title_heading_toggle(self):
        folder = self.root / "titles"
        plain = self.root / "không tiêu đề.md"
        plain.write_text("Chỉ là một đoạn.\n", encoding="utf-8")

        with_title = convert_file(plain, folder, ConversionOptions(add_title_heading=True))
        without = convert_file(plain, folder, ConversionOptions(add_title_heading=False))

        text, _ = docx_to_markdown(
            with_title.output, ConversionOptions(add_title_heading=False)
        )
        self.assertTrue(text.startswith("# không tiêu đề"))
        text, _ = docx_to_markdown(
            without.output, ConversionOptions(add_title_heading=False)
        )
        self.assertTrue(text.startswith("Chỉ là một đoạn."))

    def test_collect_files_picks_up_markdown(self):
        found = [path.name for path in collect_files([self.root])]
        self.assertIn("ghi chú.md", found)


if __name__ == "__main__":
    unittest.main(verbosity=2)
