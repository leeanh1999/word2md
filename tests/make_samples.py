"""Generate sample .docx / .xlsx files used for manual and automated testing."""

from __future__ import annotations

import io
import struct
import zipfile
import zlib
from pathlib import Path

import pandas as pd
from docx import Document
from docx.shared import Pt

SAMPLES_DIR = Path(__file__).resolve().parent / "samples"


def make_docx(path: Path) -> Path:
    document = Document()

    document.add_heading("Báo cáo kỹ thuật quý III", level=1)
    document.add_paragraph(
        "Tài liệu này được sinh tự động để kiểm thử bộ chuyển đổi Word sang Markdown."
    )

    document.add_heading("1. Tổng quan", level=2)
    paragraph = document.add_paragraph("Đoạn văn có ")
    paragraph.add_run("chữ in đậm").bold = True
    paragraph.add_run(", ")
    paragraph.add_run("chữ in nghiêng").italic = True
    paragraph.add_run(" và văn bản thường. Ký tự đặc biệt: & < > \" ' | 100%.")

    document.add_heading("2. Danh sách công việc", level=2)
    for item in ["Thu thập yêu cầu", "Thiết kế kiến trúc", "Triển khai & kiểm thử"]:
        document.add_paragraph(item, style="List Bullet")

    document.add_heading("3. Các bước thực hiện", level=2)
    for step in ["Chuẩn bị môi trường", "Chạy pipeline", "Bàn giao kết quả"]:
        document.add_paragraph(step, style="List Number")

    document.add_heading("4. Bảng số liệu", level=2)
    table = document.add_table(rows=1, cols=4)
    table.style = "Table Grid"
    headers = ["Hạng mục", "Kế hoạch", "Thực hiện", "Ghi chú"]
    for cell, text in zip(table.rows[0].cells, headers):
        run = cell.paragraphs[0].add_run(text)
        run.bold = True
        run.font.size = Pt(11)

    rows = [
        ("Doanh thu", "1.000", "1.250", "Vượt 25%"),
        ("Chi phí", "800", "770", "Tiết kiệm"),
        ("Lợi nhuận", "200", "480", "Tốt | ổn định"),
    ]
    for row in rows:
        cells = table.add_row().cells
        for cell, text in zip(cells, row):
            cell.text = text

    document.add_heading("5. Trích dẫn", level=2)
    document.add_paragraph(
        "Chất lượng không phải là một hành động, đó là một thói quen.",
        style="Intense Quote",
    )

    # A second top-level branch with three heading levels, so the navigation
    # outline is a real tree rather than a flat list.
    document.add_heading("Phụ lục", level=1)
    document.add_paragraph("Phần tham chiếu bổ sung.")

    document.add_heading("A. Thuật ngữ", level=2)
    document.add_paragraph("Định nghĩa các thuật ngữ dùng trong báo cáo.")

    document.add_heading("A.1 Viết tắt", level=3)
    document.add_paragraph("KPI, SLA, OKR.")

    document.add_heading("A.2 Tham chiếu", level=3)
    document.add_paragraph("Xem thêm tài liệu nội bộ số 42.")

    document.add_heading("B. Liên hệ", level=2)
    document.add_paragraph("Phòng Kỹ thuật — ext. 1234.")

    path.parent.mkdir(parents=True, exist_ok=True)
    document.save(path)
    return path


def make_xlsx(path: Path) -> Path:
    sales = pd.DataFrame(
        {
            "Tháng": ["01/2026", "02/2026", "03/2026"],
            "Doanh thu": [1500000, 1820000, 2100500],
            "Số đơn": [120, 143, 165],
            "Tăng trưởng": [0.0, 0.213, 0.154],
        }
    )
    staff = pd.DataFrame(
        {
            "Mã NV": ["NV001", "NV002", "NV003"],
            "Họ tên": ["Trần Văn A", "Lê Thị B", "Phạm Minh C"],
            "Ngày vào làm": pd.to_datetime(["2021-03-01", "2022-07-15", "2024-01-08"]),
            "Đang làm việc": [True, True, False],
            "Ghi chú": ["Tốt", None, "Ký tự | đặc biệt"],
        }
    )
    empty = pd.DataFrame()

    path.parent.mkdir(parents=True, exist_ok=True)
    with pd.ExcelWriter(path, engine="openpyxl") as writer:
        sales.to_excel(writer, sheet_name="Doanh thu", index=False)
        staff.to_excel(writer, sheet_name="Nhân sự", index=False)
        empty.to_excel(writer, sheet_name="Trống", index=False)
    return path


# --------------------------------------------------------- OLE attachments

# Word wraps an attached file in an OLE container and only shows an .emf icon
# for it. Building one by hand keeps the attachment tests running on machines
# without Microsoft Office.

_PACKAGE_CLSID = bytes.fromhex("0C000300" + "0000" + "0000" + "C000000000000046")

ATTACHMENTS = {
    "báo cáo.pdf": b"%PDF-1.4\n% ".ljust(4200, b"-") + b"\n%%EOF\n",
    "trang web.html": (
        "<html><body><h1>Xin chào</h1><p>".encode("utf-8")
        + b"x" * 4200
        + b"</p></body></html>"
    ),
}


def _ole10_native(label: str, data: bytes) -> bytes:
    """The `\\x01Ole10Native` stream Word writes for a packaged file."""
    source = f"D:\\nguon\\{label}"
    command = f"C:\\\\Temp\\\\{label}".encode("cp1252", "replace") + b"\0"

    def utf16(text: str) -> bytes:
        return struct.pack("<I", len(text)) + text.encode("utf-16-le")

    body = (
        struct.pack("<H", 2)
        + label.encode("cp1252", "replace")
        + b"\0"
        + source.encode("cp1252", "replace")
        + b"\0"
        + struct.pack("<HH", 0, 3)
        + struct.pack("<I", len(command))
        + command
        + struct.pack("<I", len(data))
        + data
        + utf16(f"C:\\Temp\\{label}")
        + utf16(label)
        + utf16(f"D:\\nguồn\\{label}")
    )
    return struct.pack("<I", len(body)) + body


def _compound_file(stream_name: str, payload: bytes) -> bytes:
    """A minimal v3 compound file holding one stream.

    Sector 0 is the FAT, sector 1 the directory, the rest is the payload. The
    stream has to clear the 4096-byte cut-off so it lives in normal sectors
    and no mini FAT is needed.
    """
    if len(payload) < 4096 or len(payload) > 126 * 512:
        raise ValueError("Stream phải nằm trong khoảng 4 KB - 63 KB.")

    end, free, fat_mark, none = 0xFFFFFFFE, 0xFFFFFFFF, 0xFFFFFFFD, 0xFFFFFFFF
    sectors = -(-len(payload) // 512)

    header = (
        b"\xd0\xcf\x11\xe0\xa1\xb1\x1a\xe1"
        + bytes(16)
        + struct.pack("<HHHHH", 0x003E, 3, 0xFFFE, 9, 6)
        + bytes(6)
        + struct.pack("<IIIIIIIII", 0, 1, 1, 0, 4096, end, 0, end, 0)
        + struct.pack("<I", 0)
        + struct.pack("<108I", *([free] * 108))
    )

    chain = [fat_mark, end] + [3 + index for index in range(sectors - 1)] + [end]
    fat = struct.pack("<128I", *(chain + [free] * (128 - len(chain))))

    def entry(name: str, kind: int, child: int, start: int, size: int) -> bytes:
        encoded = name.encode("utf-16-le") + b"\0\0"
        return (
            encoded.ljust(64, b"\0")
            + struct.pack("<HBB", len(encoded), kind, 1)
            + struct.pack("<III", none, none, child)
            + (_PACKAGE_CLSID if kind == 5 else bytes(16))
            + bytes(4 + 8 + 8)
            + struct.pack("<IQ", start, size)
        )

    directory = (
        entry("Root Entry", 5, 1, end, 0)
        + entry(stream_name, 2, none, 2, len(payload))
        + bytes(256)
    )

    return header + fat + directory + payload.ljust(sectors * 512, b"\0")


def _png(size: int = 8) -> bytes:
    """A solid red PNG, so the samples need no binary fixture on disk."""

    def chunk(kind: bytes, payload: bytes) -> bytes:
        return (
            struct.pack(">I", len(payload))
            + kind
            + payload
            + struct.pack(">I", zlib.crc32(kind + payload))
        )

    scanlines = b"".join(b"\0" + b"\xff\x00\x00" * size for _ in range(size))
    return (
        b"\x89PNG\r\n\x1a\n"
        + chunk(b"IHDR", struct.pack(">IIBBBBB", size, size, 8, 2, 0, 0, 0))
        + chunk(b"IDAT", zlib.compress(scanlines))
        + chunk(b"IEND", b"")
    )


def _insert_picture(document, after_heading: str, image: bytes) -> None:
    """Put a picture in its own paragraph, right under the given heading."""
    paragraphs = document.paragraphs
    index = next(i for i, p in enumerate(paragraphs) if p.text == after_heading)
    anchor = paragraphs[index + 1]
    anchor.insert_paragraph_before().add_run().add_picture(io.BytesIO(image))


def make_attachments_docx(base: Path, path: Path) -> Path:
    """Copy a .docx and graft a picture and OLE attachments onto it.

    The picture and the attachments land in different sections, so exporting
    one section can be checked to leave the other one's assets behind.
    """
    icon = b"\x01\x00\x00\x00" + bytes(84)  # enough of an .emf header to link
    paragraphs = []
    extra: dict[str, bytes] = {"word/media/icon1.emf": icon}
    relationships = [
        '<Relationship Id="rIdIcon" Type="http://schemas.openxmlformats.org/'
        'officeDocument/2006/relationships/image" Target="media/icon1.emf"/>'
    ]

    for index, (name, data) in enumerate(ATTACHMENTS.items(), start=1):
        part = f"embeddings/oleObject{index}.bin"
        extra[f"word/{part}"] = _compound_file(
            "\x01Ole10Native", _ole10_native(name, data)
        )
        relationships.append(
            f'<Relationship Id="rIdOle{index}" Type="http://schemas.openxmlformats'
            f'.org/officeDocument/2006/relationships/oleObject" Target="{part}"/>'
        )
        paragraphs.append(
            "<w:p><w:r>"
            '<w:object w:dxaOrig="1535" w:dyaOrig="998">'
            f'<v:shape id="_x0000_i{1024 + index}" type="#_x0000_t75" o:ole="">'
            '<v:imagedata r:id="rIdIcon" o:title=""/></v:shape>'
            f'<o:OLEObject Type="Embed" ProgID="Package" DrawAspect="Icon" '
            f'ObjectID="_{index}" r:id="rIdOle{index}"/>'
            "</w:object></w:r>"
            f'<w:r><w:t xml:space="preserve"> là tệp đính kèm.</w:t></w:r></w:p>'
        )

    illustrated = io.BytesIO()
    document = Document(str(base))
    _insert_picture(document, "A.1 Viết tắt", _png())
    document.save(illustrated)

    path.parent.mkdir(parents=True, exist_ok=True)
    with zipfile.ZipFile(illustrated) as source, zipfile.ZipFile(
        path, "w", zipfile.ZIP_DEFLATED
    ) as target:
        for item in source.infolist():
            content = source.read(item.filename)
            if item.filename == "word/document.xml":
                text = content.decode("utf-8").replace(
                    "</w:body>", "".join(paragraphs) + "</w:body>"
                )
                content = text.encode("utf-8")
            elif item.filename == "word/_rels/document.xml.rels":
                text = content.decode("utf-8").replace(
                    "</Relationships>", "".join(relationships) + "</Relationships>"
                )
                content = text.encode("utf-8")
            elif item.filename == "[Content_Types].xml":
                text = content.decode("utf-8").replace(
                    "</Types>",
                    '<Default Extension="emf" ContentType="image/x-emf"/>'
                    '<Default Extension="bin" ContentType="application/'
                    'vnd.openxmlformats-officedocument.oleObject"/></Types>',
                )
                content = text.encode("utf-8")
            target.writestr(item, content)
        for name, blob in extra.items():
            target.writestr(name, blob)
    return path


def make_corrupt(path: Path) -> Path:
    """A file with a valid extension but garbage content, to test error handling."""
    path.parent.mkdir(parents=True, exist_ok=True)
    path.write_bytes(b"this is definitely not a zip container")
    return path


def make_mislabelled(source: Path, path: Path) -> Path:
    """A modern file wearing a legacy extension - very common in the wild."""
    path.parent.mkdir(parents=True, exist_ok=True)
    path.write_bytes(source.read_bytes())
    return path


def make_legacy(docx: Path, xlsx: Path, directory: Path) -> dict[str, Path]:
    """Downgrade the samples to .doc/.xls via Office COM, when available.

    Returns an empty dict on machines without Microsoft Office; the tests that
    need real binary samples skip themselves in that case.
    """
    import sys

    sys.path.insert(0, str(Path(__file__).resolve().parents[1]))
    from src.legacy import LegacyUpgrader, has_msoffice

    if not (has_msoffice("Word.Application") and has_msoffice("Excel.Application")):
        return {}

    WD_FORMAT_97 = 0
    XL_FORMAT_97 = 56
    created: dict[str, Path] = {}
    upgrader = LegacyUpgrader()
    try:
        upgrader._initialize_com()

        word = upgrader._word_app()
        document = word.Documents.Open(str(docx.resolve()), ReadOnly=True, Visible=False)
        try:
            target = directory / "test.doc"
            document.SaveAs2(str(target.resolve()), FileFormat=WD_FORMAT_97)
            created["doc"] = target
        finally:
            document.Close(0)

        excel = upgrader._excel_app()
        workbook = excel.Workbooks.Open(str(xlsx.resolve()), ReadOnly=True)
        try:
            target = directory / "test.xls"
            workbook.SaveAs(str(target.resolve()), FileFormat=XL_FORMAT_97)
            created["xls"] = target
        finally:
            workbook.Close(False)
    finally:
        upgrader.close()

    return created


def build_all(directory: Path = SAMPLES_DIR, legacy: bool = True) -> dict[str, Path]:
    docx = make_docx(directory / "test.docx")
    xlsx = make_xlsx(directory / "test.xlsx")
    samples = {
        "docx": docx,
        "xlsx": xlsx,
        "attachments": make_attachments_docx(docx, directory / "attachments.docx"),
        "corrupt": make_corrupt(directory / "corrupt.docx"),
        "fake_doc": make_mislabelled(docx, directory / "mislabelled.doc"),
        "fake_xls": make_mislabelled(xlsx, directory / "mislabelled.xls"),
    }
    if legacy:
        samples.update(make_legacy(docx, xlsx, directory))
    return samples


if __name__ == "__main__":
    for kind, created in build_all().items():
        print(f"{kind:10} -> {created}")
